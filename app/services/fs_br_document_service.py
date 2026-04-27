
import io
import os
from typing import List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# â”€â”€â”€ Paths & constants â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..", "static", "docs",
    "Functional_Specification_Template_en.docx",
)

BR_TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..", "static", "docs",
    "Business_requirement_template.docx",
)

FS_VARIANT_TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..", "static", "docs",
    "FS Template.docx",
)

# XML namespace for Word 2010 checkbox SDT elements
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"

BSH_GREY = RGBColor(0x64, 0x74, 0x8B)


# â”€â”€â”€ Low-level helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def _set_para_text(para, text: str):
    """Replace a paragraph's text, preserving the first run's formatting."""
    if para.runs:
        for run in para.runs[1:]:
            run.text = ""
        para.runs[0].text = text
    else:
        para.add_run(text)


def _set_table_cell(table, row_idx, col_idx, text: str):
    """Write text into a table cell, preserving existing run formatting."""
    cell = table.cell(row_idx, col_idx)
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)


def _find_heading_para(doc, heading_text: str):
    """Find a paragraph whose text matches heading_text (case-insensitive prefix)."""
    target = heading_text.strip().lower()
    for p in doc.paragraphs:
        if "Heading" in p.style.name and p.text.strip().lower().startswith(target):
            return p
    return None


def _first_content_para_after(doc, heading_para):
    """Return the first non-empty Normal paragraph after a heading."""
    found = False
    for p in doc.paragraphs:
        if p._p is heading_para._p:
            found = True
            continue
        if found:
            if "Heading" in p.style.name:
                return None  # next heading reached, no content
            if p.text.strip():
                return p
    return None


def _get_content_paras_between_headings(doc, heading_para):
    """Return ALL content paragraphs (empty or not) between heading_para and the next heading."""
    paras = []
    found = False
    for p in doc.paragraphs:
        if p._p is heading_para._p:
            found = True
            continue
        if found:
            if "Heading" in p.style.name:
                break
            paras.append(p)
    return paras


def _remove_paragraph(para):
    """Remove a paragraph element from the document body."""
    p_elem = para._element
    parent = p_elem.getparent()
    if parent is not None:
        parent.remove(p_elem)


def _clear_section_set_text(doc, heading_para, text: str):
    """Remove ALL content paragraphs under a heading, then set text on the first one.
    If text is empty, leave the section cleared with one blank paragraph.
    """
    content_paras = _get_content_paras_between_headings(doc, heading_para)
    if not content_paras:
        # No content paras â€” insert one
        if text:
            new_p = OxmlElement("w:p")
            new_r = OxmlElement("w:r")
            new_t = OxmlElement("w:t")
            new_t.set(qn("xml:space"), "preserve")
            new_t.text = text
            new_r.append(new_t)
            new_p.append(new_r)
            heading_para._p.addnext(new_p)
        return

    # Set first paragraph to user text (or blank)
    _set_para_text(content_paras[0], text if text else "")

    # Remove all remaining content paragraphs
    for p in content_paras[1:]:
        _remove_paragraph(p)


def _ensure_table_rows(table, needed_data_rows: int, header_rows: int = 1):
    """Expand a Word table so it has at least `needed_data_rows` data rows
    (below `header_rows` header rows). Rows are cloned from the last existing
    data row to preserve formatting. Existing rows are NOT removed."""
    import copy
    current_total = len(table.rows)
    current_data = current_total - header_rows
    if needed_data_rows <= current_data:
        return
    last_row_elem = table.rows[-1]._tr
    for _ in range(needed_data_rows - current_data):
        new_tr = copy.deepcopy(last_row_elem)
        for cell_el in new_tr.findall(qn("w:tc")):
            for p_el in cell_el.findall(qn("w:p")):
                for r_el in p_el.findall(qn("w:r")):
                    for t_el in r_el.findall(qn("w:t")):
                        t_el.text = ""
        table._tbl.append(new_tr)


def _fill_table(table, data_rows, num_cols, header_rows=1):
    """Expand a table if needed, then fill it with the provided data rows."""
    _ensure_table_rows(table, len(data_rows), header_rows)
    for i, row_data in enumerate(data_rows):
        row_idx = i + header_rows
        for col_idx in range(min(len(row_data), num_cols)):
            val = row_data[col_idx].strip() if row_data[col_idx] else ""
            if val:
                _set_table_cell(table, row_idx, col_idx, val)


# â”€â”€â”€ Functional Specification generator â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€


def generate_functional_spec_docx(data: dict) -> io.BytesIO:
    """
    Fill the BSH Functional Specification template with form data.
    Returns an in-memory BytesIO buffer ready for streaming.

    Template structure (Functional_Specification_Template_en.docx):
      Page Header Table (2x3):
        R0C2: Date / Version / Author  |  R1: Title (merged)
      Body Tables:
        Table 0 (5x4): Responsibilities - Function/Name/Date/Task
        Table 1 (4x4): Previous Steps  - Step/When/Where/Who
        Table 2 (8x2): Glossary        - Abbreviation/Description
        Table 3 (4x5): Document History - Version/Date/Actioned by/Description/Status
      Sections (Heading paragraphs):
        1.  Starting point
        1.1 (Related) project goal      -> P54
        1.2 Developer statement          (boilerplate, kept)
        2.  Solution description         -> P66
        2.1 Improvement potential        -> P70
        2.2 Delimitation of solution     -> P74
        2.3 Previous steps               -> Table 1
        3.  Solution definition
        3.1 Functionality                -> P90
        3.2 User view, dialog execution  -> P94
        3.4 Data structures              -> P111
        3.5 Data maintenance             -> P115
        3.6 Interfaces                   -> P119
        3.7 Authorization â€“ user roles   -> P123
        3.8 Information security         -> (after P130)
        3.10 Architecture and technology -> P161
        5.  Risks                        -> P175
        6.  List of open issues          -> P179
        7.  Migration                    -> P183
        8.  Glossary                     -> Table 2
        9.  Document history             -> Table 3
    """
    doc = Document(TEMPLATE_PATH)
    tables = doc.tables


    # 1. PAGE HEADER â€” Title, Date, Version, Author

    title = data.get("title", "").strip()
    date_val = data.get("date", "").strip()
    version = data.get("version", "").strip()
    author = data.get("author", "").strip()

    hdr_table = doc.sections[0].header.tables[0]

    # Row 0, Col 2 has runs: "Date" \t "xx.xx.xxxx" \n "Version" \t "1.0" \n "Author" \t
    cell_r0c2 = hdr_table.cell(0, 2)
    runs = cell_r0c2.paragraphs[0].runs
    if len(runs) >= 7:
        runs[2].text = date_val or "xx.xx.xxxx"   # date value
        runs[4].text = f"\t{version}" if version else "\t1.0"  # version value
        runs[6].text = f"\t{author}"    # author value

    # Row 1 (merged): Title
    if title:
        for ci in range(3):
            cell = hdr_table.cell(1, ci)
            para = cell.paragraphs[0]
            if para.runs:
                para.runs[0].text = title
                for r in para.runs[1:]:
                    r.text = ""


    # 2. RESPONSIBILITIES TABLE (Table 0)

    resp_table = tables[0]
    responsibilities = data.get("responsibilities", {})

    resp_map = {
        1: "globalBusiness",
        2: "globalShape",
        3: "developer",
        4: "steward",
    }

    for row_idx, key in resp_map.items():
        entry = responsibilities.get(key, {})
        name = entry.get("name", "").strip()
        date_str = entry.get("date", "").strip()
        if name:
            _set_table_cell(resp_table, row_idx, 1, name)
        if date_str:
            _set_table_cell(resp_table, row_idx, 2, date_str)


    # 3. SECTION CONTENT â€” Replace placeholder text after headings


    # Sections that need ALL content paragraphs cleared (multi-paragraph boilerplate)
    # These sections have boilerplate text that must be fully replaced by user input.
    multi_para_sections = [
        ("User view", "userView"),
        ("Language topics", "languageTopics"),
        ("Authorization", "authorization"),
        ("Migration", "migration"),
    ]

    for heading_text, data_key in multi_para_sections:
        user_text = data.get(data_key, "").strip()
        heading_para = _find_heading_para(doc, heading_text)
        if not heading_para:
            continue
        _clear_section_set_text(doc, heading_para, user_text)

    # â”€â”€ Solution definition intro (Report / Transaction / Source System) â”€â”€
    # P66-70: "For detailed analysis..." + Report + Transaction + Source system
    # Clear boilerplate, replace with user's report/transaction/sourceSystem
    sol_def_heading = _find_heading_para(doc, "Solution definition")
    if sol_def_heading:
        report = data.get("report", "").strip()
        transaction = data.get("transaction", "").strip()
        source_system = data.get("sourceSystem", "").strip()

        content_paras = _get_content_paras_between_headings(doc, sol_def_heading)
        if content_paras:
            # First para (P66): "For detailed analysis..." â†’ set Report value
            if report:
                _set_para_text(content_paras[0], f"Report:\t\t{report}")
            else:
                _set_para_text(content_paras[0], "")

            # Remaining paragraphs: set Transaction, Source System, clear the rest
            idx = 1
            if idx < len(content_paras):
                # Was P68 (Report): now becomes Transaction
                if transaction:
                    _set_para_text(content_paras[idx], f"Transaction:\t\t{transaction}")
                else:
                    _set_para_text(content_paras[idx], "")
                idx += 1
            if idx < len(content_paras):
                # Was P69 (Transaction): now becomes Source system
                if source_system:
                    _set_para_text(content_paras[idx], f"Source system:\t{source_system}")
                else:
                    _set_para_text(content_paras[idx], "")
                idx += 1
            # Remove any remaining content paras
            for p in content_paras[idx:]:
                _remove_paragraph(p)

    # Sections that only need the first content paragraph replaced (simple ones)
    simple_sections = [
        ("(Related) project goal", "projectGoal"),
        ("Solution description", "solutionDesc"),
        ("Improvement potential", "improvementPotential"),
        ("Delimitation of solution", "delimitation"),
        ("Functionality", "functionality"),
        ("Data structures", "dataStructures"),
        ("Data maintenance", "dataMaintenance"),
        ("Interfaces", "interfaces"),
        ("Information security", "infoSecurity"),
        ("Architecture and technology", "architecture"),
        ("Risks", "risks"),
        ("List of open issues", "openIssues"),
    ]

    for heading_text, data_key in simple_sections:
        user_text = data.get(data_key, "").strip()
        if not user_text:
            continue
        heading_para = _find_heading_para(doc, heading_text)
        if not heading_para:
            continue
        content_para = _first_content_para_after(doc, heading_para)
        if content_para:
            _set_para_text(content_para, user_text)
        else:
            new_p = OxmlElement("w:p")
            new_r = OxmlElement("w:r")
            new_t = OxmlElement("w:t")
            new_t.set(qn("xml:space"), "preserve")
            new_t.text = user_text
            new_r.append(new_t)
            new_p.append(new_r)
            heading_para._p.addnext(new_p)


    # 4. PREVIOUS STEPS TABLE (Table 1)

    prev_steps = data.get("previousSteps", [])
    if prev_steps:
        _fill_table(tables[1], prev_steps, 4)


    # 5. GLOSSARY TABLE (Table 2)

    glossary = data.get("glossary", [])
    if glossary:
        _fill_table(tables[2], glossary, 2)


    # 6. DOCUMENT HISTORY TABLE (Table 3)

    doc_history = data.get("docHistory", [])
    if doc_history:
        _fill_table(tables[3], doc_history, 5)


    # 7. FOOTER â€” update title in footer

    if title:
        footer = doc.sections[0].footer
        if footer and not footer.is_linked_to_previous:
            for p in footer.paragraphs:
                if "Business Requirement" in p.text or "XXXX" in p.text:
                    _set_para_text(p, f"\t\t\n{title}\tPage 1\tDate of printing: ")

    # -- Write to buffer --
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# â”€â”€â”€ Business Requirement generator â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def _set_table_cell(table, row_idx, col_idx, text: str):
    """Set text in a specific table cell, preserving first run formatting."""
    cell = table.rows[row_idx].cells[col_idx]
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.text = text


def _append_text_after_heading(doc, heading_para, text: str):
    """Insert a new paragraph with text right after a heading paragraph."""
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")

    # Match document default font
    rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")  # 11pt
    rpr.append(sz)
    new_r.append(rpr)

    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = text
    new_r.append(new_t)
    new_p.append(new_r)

    heading_para._element.addnext(new_p)
    return new_p


def generate_br_docx(data: dict) -> io.BytesIO:
    """
    Fill the BSH Business Requirement template with form data.
    Returns an in-memory BytesIO buffer ready for streaming.

    Template structure (Business_requirement_template.docx):
      Paragraphs:
        [ 0] "* indicates a mandatory field"
        [ 5] Heading 1 "Business Requirement Description *"
        [ 6] Heading 2 "What is the initial situation?"
        [ 7] Heading 2 "What is the required situation?"
        [ 8] Heading 2 "Who is involved in the function/process (departments)?"
        [ 9] Heading 2 "Which IT-components are used â€¦?"
        [10] Heading 2 "What data is used or affected?"
        [11] Heading 2 "What is your proposal for solution?"
        [13] Heading 1 "Benefits for the Business"
        [14] Heading 2 "What benefits can be reached?"
        [15] Heading 2 "What are your savings p.a.?"
        [16] Heading 2 "What would happen without the implementation?"

      Tables:
        Table 0 (4x4): Header info (title, project, product owner, IT product, target date, requestor, company)
        Table 2 (5x4): Responsibles (Local Business, Global Business, Global Shape GDS, Regional Shape GDS)
    """
    doc = Document(BR_TEMPLATE_PATH)
    paras = doc.paragraphs
    tables = doc.tables


    # 1. HEADER TABLE (Table 0)

    header_table = tables[0]

    # Row 0: Requirement Title (merged across all 4 cols â†’ set col 0)
    title = data.get("title", "").strip()
    if title:
        # The title row cells are merged; set the first cell
        _set_table_cell(header_table, 0, 0, title)

    # Row 1: Project (col 1), Product Owner (col 3)
    project = data.get("project", "").strip()
    product_owner = data.get("productOwner", "").strip()
    if project:
        _set_table_cell(header_table, 1, 1, project)
    if product_owner:
        _set_table_cell(header_table, 1, 3, product_owner)

    # Row 2: IT Product (col 1), Target Date (col 3)
    it_product = data.get("itProduct", "").strip()
    target_date = data.get("targetDate", "").strip()
    if it_product:
        _set_table_cell(header_table, 2, 1, it_product)
    if target_date:
        _set_table_cell(header_table, 2, 3, target_date)

    # Row 3: Requestor (col 1), Requestor Company (col 3)
    requestor = data.get("requestor", "").strip()
    requestor_company = data.get("requestorCompany", "").strip()
    if requestor:
        _set_table_cell(header_table, 3, 1, requestor)
    if requestor_company:
        _set_table_cell(header_table, 3, 3, requestor_company)


    # DOCUMENT HEADER (page header) â€“ Create Date & Created By

    create_date = data.get("createDate", "").strip()
    created_by = data.get("createdBy", "").strip()
    if create_date or created_by:
        hdr_table = doc.sections[0].header.tables[0]
        cell = hdr_table.cell(0, 2)
        para = cell.paragraphs[0]
        # The cell has 2 runs: "Create Date: " and "\nCreated by: "
        runs = para.runs
        if len(runs) >= 2:
            runs[0].text = f"Create Date: {create_date}"
            runs[1].text = f"\nCreated by: {created_by}"


    # 2. RESPONSIBLES TABLE (Table 2)

    resp_table = tables[2]
    responsibles = data.get("responsibles", {})

    # Row mapping: 1=Local Business, 2=Global Business, 3=Global Shape GDS, 4=Regional Shape GDS
    # Columns: 0=Role, 1=eMail, 2=Company, 3=Department
    resp_map = {
        1: "localBusiness",
        2: "globalBusiness",
        3: "globalShapeGds",
        4: "regionalShapeGds",
    }

    for row_idx, key in resp_map.items():
        entry = responsibles.get(key, {})
        email = entry.get("email", "").strip()
        company = entry.get("company", "").strip()
        dept = entry.get("department", "").strip()
        if email:
            _set_table_cell(resp_table, row_idx, 1, email)
        if company:
            _set_table_cell(resp_table, row_idx, 2, company)
        if dept:
            _set_table_cell(resp_table, row_idx, 3, dept)


    # 3. BUSINESS REQUIREMENT DESCRIPTION (paras 6-11)

    description = data.get("description", {})

    field_map = {
        6: "initialSituation",
        7: "requiredSituation",
        8: "departments",
        9: "itComponents",
        10: "dataAffected",
        11: "proposalSolution",
    }

    for para_idx, field_key in field_map.items():
        text = description.get(field_key, "").strip()
        if text:
            _append_text_after_heading(doc, paras[para_idx], text)


    # 4. BENEFITS (paras 14-16)

    benefits = data.get("benefits", {})

    benefits_map = {
        14: "benefitsReached",
        15: "savingsPA",
        16: "withoutImplementation",
    }

    for para_idx, field_key in benefits_map.items():
        text = benefits.get(field_key, "").strip()
        if text:
            _append_text_after_heading(doc, paras[para_idx], text)


    # 5. SIGN OFF / AGREEMENT (Table 3)
    #    Row 0: header (Agreement, Name, Department, Date, Signature)
    #    Row 1: Global Business Process Owner
    #    Row 2: GDS Product Owner / Manager

    sign_off = data.get("signOff", {})
    signoff_table = tables[3]

    gbpo = sign_off.get("gbpo", {})
    if gbpo.get("name", "").strip():
        _set_table_cell(signoff_table, 1, 1, gbpo["name"].strip())
    if gbpo.get("department", "").strip():
        _set_table_cell(signoff_table, 1, 2, gbpo["department"].strip())
    if gbpo.get("date", "").strip():
        _set_table_cell(signoff_table, 1, 3, gbpo["date"].strip())

    gds = sign_off.get("gds", {})
    if gds.get("name", "").strip():
        _set_table_cell(signoff_table, 2, 1, gds["name"].strip())
    if gds.get("department", "").strip():
        _set_table_cell(signoff_table, 2, 2, gds["department"].strip())
    if gds.get("date", "").strip():
        _set_table_cell(signoff_table, 2, 3, gds["date"].strip())


    # 6. COST ESTIMATION & DECISION

    decision = data.get("decision", {})

    # Table 5: Evaluation (single-cell table)
    evaluation = decision.get("evaluation", "").strip()
    if evaluation:
        cell = tables[5].rows[0].cells[0]
        existing = cell.text.strip()
        cell.paragraphs[0].text = existing
        # Add answer on new line
        p = cell.add_paragraph()
        run = p.add_run(evaluation)
        run.font.size = Pt(11)

    # Table 6: Decision date (single-cell table)
    decision_date = decision.get("decisionDate", "").strip()
    if decision_date:
        cell = tables[6].rows[0].cells[0]
        existing = cell.text.strip()
        cell.paragraphs[0].text = f"{existing} {decision_date}"

    # Table 7: Decision type â€” custom bigger checkboxes in column 0
    decision_type = decision.get("decisionType", "").strip()
    accepted_sub = decision.get("acceptedSubOptions", [])

    # Table structure:
    #   Row 0: rejected           (main option)
    #   Row 1: accepted           (main option)
    #   Row 2: add-on to BSH     (sub-option of accepted)
    #   Row 3: only local         (sub-option of accepted)
    #   Row 4: accepted w/ restr  (main option)
    #
    # For all 5 rows: put custom checkbox in col 0 (â˜ or â˜’)
    decision_table = tables[7]
    decision_check_map = {
        0: decision_type == "rejected",
        1: decision_type == "accepted",
        2: decision_type == "accepted" and "add-on" in accepted_sub,
        3: decision_type == "accepted" and "local-only" in accepted_sub,
        4: decision_type == "accepted-with-restrictions",
    }
    for row_idx in range(5):
        checked = decision_check_map.get(row_idx, False)
        _set_custom_checkbox(decision_table, row_idx, 0, checked)

    # Table 8: Specification of restrictions (single-cell table)
    restrictions = decision.get("restrictions", "").strip()
    if restrictions:
        cell = tables[8].rows[0].cells[0]
        existing = cell.text.strip()
        cell.paragraphs[0].text = existing
        p = cell.add_paragraph()
        run = p.add_run(restrictions)
        run.font.size = Pt(11)

    # Table 9: Reason for decision (single-cell table)
    reason = decision.get("reason", "").strip()
    if reason:
        cell = tables[9].rows[0].cells[0]
        existing = cell.text.strip()
        cell.paragraphs[0].text = existing
        p = cell.add_paragraph()
        run = p.add_run(reason)
        run.font.size = Pt(11)

    # Table 10: Target date in case of implementation (single-cell table)
    impl_target_date = decision.get("implTargetDate", "").strip()
    if impl_target_date:
        cell = tables[10].rows[0].cells[0]
        existing = cell.text.strip()
        cell.paragraphs[0].text = f"{existing} {impl_target_date}"


    # 7. COSTS / SAVINGS / CHARGING (Table 12)
    #    Row 0: headers
    #    Rows 1-3: data rows
    #    Row 4: totals

    costs = data.get("costs", {})
    cost_rows = costs.get("rows", [])
    cost_table = tables[12]

    # Expand table if more than 3 data rows (totals row will shift down)
    needed_data = max(len(cost_rows), 3)
    _ensure_table_rows(cost_table, needed_data + 1, header_rows=1)  # +1 for totals row
    totals_row_idx = needed_data + 1  # row after last data row

    total_initial = 0
    total_running = 0
    total_savings = 0
    total_payback = 0
    period_sum = 0
    period_count = 0

    for i, cr in enumerate(cost_rows):
        row_idx = i + 1  # data rows start at 1
        it_prod = cr.get("itProduct", "").strip()
        catsnr = cr.get("catsnr", "").strip()
        initial = cr.get("initialCosts", "").strip()
        running = cr.get("runningCosts", "").strip()
        savings_val = cr.get("savings", "").strip()

        if it_prod:
            _set_table_cell(cost_table, row_idx, 0, it_prod)
        if catsnr:
            _set_table_cell(cost_table, row_idx, 1, catsnr)

        initial_f = float(initial) if initial else 0
        running_f = float(running) if running else 0
        savings_f = float(savings_val) if savings_val else 0

        if initial:
            _set_table_cell(cost_table, row_idx, 2, initial)
        if running:
            _set_table_cell(cost_table, row_idx, 3, running)
        if savings_val:
            _set_table_cell(cost_table, row_idx, 4, savings_val)

        payback = savings_f - running_f
        _set_table_cell(cost_table, row_idx, 5, f"{payback:.0f}" if (initial or running or savings_val) else "-")

        if payback > 0:
            period = initial_f / payback
            _set_table_cell(cost_table, row_idx, 6, f"{period:.1f}")
            period_sum += period
            period_count += 1
        else:
            _set_table_cell(cost_table, row_idx, 6, "-")

        total_initial += initial_f
        total_running += running_f
        total_savings += savings_f
        total_payback += payback

    # Fill totals row
    has_data = total_initial or total_running or total_savings
    if has_data:
        _set_table_cell(cost_table, totals_row_idx, 2, f"{total_initial:.0f}")
        _set_table_cell(cost_table, totals_row_idx, 3, f"{total_running:.0f}")
        _set_table_cell(cost_table, totals_row_idx, 4, f"{total_savings:.0f}")
        _set_table_cell(cost_table, totals_row_idx, 5, f"{total_payback:.0f}")
        _set_table_cell(cost_table, totals_row_idx, 6, f"{period_sum / period_count:.1f}" if period_count else "-")

    # -- Write to buffer --
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _set_custom_checkbox(table, row_idx, col_idx, checked: bool):
    """
    Replace the content of a cell with a large custom checkbox character.
    â˜’ (checked) or â˜ (unchecked), rendered at 16pt for visibility.
    Clears any existing content/borders in the cell first.
    """
    char = "\u2612" if checked else "\u2610"
    cell = table.rows[row_idx].cells[col_idx]

    # Clear all existing paragraphs
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""

    # Set the checkbox character in the first paragraph
    para = cell.paragraphs[0]
    if para.runs:
        run = para.runs[0]
    else:
        run = para.add_run()

    run.text = char
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Center the character
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# â”€â”€â”€ FS Template (Variant) generator â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def generate_fs_variant_docx(data: dict) -> io.BytesIO:
    """
    Fill the FS Template (variant) with form data.
    Returns an in-memory BytesIO buffer ready for streaming.

    Template structure (FS Template.docx):
      Cover page paragraphs:
        [4]  Description:
        [10] Written By:
        [11] Date:
      Table 0 (6x8): Revision History
      Â§1  Purpose (para 55) + Type / Latency / Frequency / System / Impacted System
      Â§2  Detail Processing Logic (after heading at para 66)
          Prerequisites/Assumptions (para 68)
      Â§3  Program Inputs
        3.1 Selection Screen  â€” Table 1 (3 rows x 6 cols)
        3.2 Report Characteristic â€” Table 2 (5 rows x 3 cols)
        3.3 Report Delivery (para 84-85)
        3.4 Report/Form Layout (after heading at para 86)
        3.5 Report Attributes (para 89)
        3.6 Custom Transitions (para 92)
        3.7 Printer Requirements (para 95)
        3.8 Exclusions (para 97) + Output Files
        3.9 Exception Handling (para 103)
        3.11 Constraints (para 105)
        3.10 Dependencies (para 107)
        3.12 Scheduling Requirements (para 109)
        3.13 Role/Authorization (para 111)
        3.14 Test Specification â€” Table 5 (5x4), Table 6 (2x2)
      Â§4  Change History â€” Table 7 (7x4)
    """
    doc = Document(FS_VARIANT_TEMPLATE_PATH)
    paras = doc.paragraphs
    tables = doc.tables


    # 1. COVER PAGE â€” Description, Written By, Date

    description_text = data.get("description", "").strip()
    written_by = data.get("writtenBy", "").strip()
    date_val = data.get("date", "").strip()

    if description_text and len(paras) > 4:
        _set_para_text(paras[4], f"Description: {description_text}")
    if written_by and len(paras) > 10:
        _set_para_text(paras[10], f"Written By: {written_by}")
    if date_val and len(paras) > 11:
        _set_para_text(paras[11], f"Date: {date_val}")

    # Footer â€” "Updated By : â€¦" and "Version : â€¦"
    updated_by = data.get("updatedBy", "").strip()
    version_val = data.get("version", "").strip()
    if updated_by or version_val:
        for section in doc.sections:
            footer = section.footer
            if footer:
                for ft in footer.tables:
                    if len(ft.rows) >= 1 and len(ft.columns) >= 3:
                        if updated_by:
                            _set_table_cell(ft, 0, 0, f"Updated By : {updated_by}")
                        if version_val:
                            _set_table_cell(ft, 0, 2, f"Version : {version_val}")


    # 2. REVISION HISTORY TABLE (Table 0)
    #    Cols: Version, Effective Date, Brief Description, Reference,
    #          Affected Sections, Prepared By, Reviewed By, Approved By

    revision_history = data.get("revisionHistory", [])
    if revision_history:
        _fill_table(tables[0], revision_history, 8)


    # 3. PURPOSE SECTION

    purpose = data.get("purpose", "").strip()
    if purpose and len(paras) > 55:
        _set_para_text(paras[55], purpose)

    # Type checkboxes (para 57)
    type_val = data.get("type", "").strip()
    if type_val and len(paras) > 57:
        _set_para_text(paras[57], f"Type:\t{type_val}")

    # Latency (para 59)
    latency = data.get("latency", "").strip()
    if latency and len(paras) > 59:
        _set_para_text(paras[59], f"Latency: {latency}")

    # Frequency (para 61)
    frequency = data.get("frequency", "").strip()
    if frequency and len(paras) > 61:
        _set_para_text(paras[61], f"Frequency: \t{frequency}")
        if len(paras) > 62:
            _set_para_text(paras[62], "")

    # System (para 64)
    system = data.get("system", "").strip()
    if system and len(paras) > 64:
        _set_para_text(paras[64], f"System\t\t{system}")

    # Impacted System (para 65)
    impacted_system = data.get("impactedSystem", "").strip()
    if impacted_system and len(paras) > 65:
        _set_para_text(paras[65], f"Impacted System:\t{impacted_system}")


    # 4. DETAIL PROCESSING LOGIC

    processing_logic = data.get("processingLogic", "").strip()
    if processing_logic:
        heading = _find_heading_para(doc, "2. Detail Processing Logic")
        if heading:
            content = _first_content_para_after(doc, heading)
            if content:
                _set_para_text(content, processing_logic)
            else:
                _append_text_after_heading(doc, heading, processing_logic)

    prerequisites = data.get("prerequisites", "").strip()
    if prerequisites and len(paras) > 69:
        # Para 68 is "Prerequisites/Assumptions" label, 69 is content
        prereq_heading = _find_heading_para(doc, "Prerequisites")
        if not prereq_heading:
            # It's a normal paragraph, find para 68 and add content after
            _set_para_text(paras[69], prerequisites)


    # 5. SELECTION SCREEN TABLE (Table 1)
    #    Cols: Screen Label, Referenced Field, Range of Value,
    #          Attributes/Defaults, Validations, Comments

    selection_screen = data.get("selectionScreen", [])
    if selection_screen:
        _fill_table(tables[1], selection_screen, 6)


    # 6. REPORT CHARACTERISTICS (Table 2)
    #    Rows: Standard, Interactive, Drill-Down, ALV, Other
    #    Cols: Type, Yes/No, Comments

    report_chars = data.get("reportCharacteristics", {})
    char_table = tables[2]
    char_map = {
        0: "standard",
        1: "interactive",
        2: "drillDown",
        3: "alv",
        4: "other",
    }
    for row_idx, key in char_map.items():
        entry = report_chars.get(key, {})
        yn = entry.get("value", "").strip()
        comment = entry.get("comments", "").strip()
        if yn:
            _set_table_cell(char_table, row_idx, 1, yn)
        if comment:
            _set_table_cell(char_table, row_idx, 2, comment)


    # 7. REPORT DELIVERY (para 84-85)

    report_delivery = data.get("reportDelivery", "").strip()
    if report_delivery and len(paras) > 84:
        _set_para_text(paras[84], report_delivery)
        if len(paras) > 85:
            _set_para_text(paras[85], "")


    # 8. SIMPLE SECTIONS (heading â†’ content para replacement)

    simple_sections = [
        ("Report/Form Layout", "reportLayout"),
        ("Report Attributes", "reportAttributes"),
        ("Custom Transitions", "customTransitions"),
        ("Printer Requirements", "printerRequirements"),
        ("Exclusions", "exclusions"),
        ("Exception Handling", "exceptionHandling"),
        ("Constraints", "constraints"),
        ("Dependencies", "dependencies"),
        ("Scheduling Requirements", "scheduling"),
        ("Role/Authorization", "roleAuthorization"),
    ]

    for heading_text, data_key in simple_sections:
        user_text = data.get(data_key, "").strip()
        if not user_text:
            continue
        heading_para = _find_heading_para(doc, heading_text)
        if not heading_para:
            continue
        content_para = _first_content_para_after(doc, heading_para)
        if content_para:
            _set_para_text(content_para, user_text)

    # Output Files section (paras 98-100, under Exclusions heading)
    output_file_location = data.get("outputFileLocation", "").strip()
    if output_file_location and len(paras) > 99:
        _set_para_text(paras[99], f"Output File Location: {output_file_location}")

    output_file_remarks = data.get("outputFileRemarks", "").strip()
    if output_file_remarks and len(paras) > 100:
        _set_para_text(paras[100], f"Additional Output File Location Remarks: {output_file_remarks}")


    # 9. TEST SPECIFICATION

    # Table 5 (5x4): Test scenarios â€” ID, Test Scenario, Expected Results, Comments
    test_scenarios = data.get("testScenarios", [])
    if test_scenarios:
        _fill_table(tables[5], test_scenarios, 4)

    # Test Data text (after "Test Data & Other Needs" heading)
    test_data = data.get("testData", "").strip()
    if test_data:
        td_heading = _find_heading_para(doc, "Test Data")
        if td_heading:
            content = _first_content_para_after(doc, td_heading)
            if content:
                _set_para_text(content, test_data)

    # Table 6 (2x2): Test System & Client
    test_system = data.get("testSystem", "").strip()
    test_client = data.get("testClient", "").strip()
    env_table = tables[6]
    if test_system:
        _set_table_cell(env_table, 0, 1, test_system)
    if test_client:
        _set_table_cell(env_table, 1, 1, test_client)


    # 10. CHANGE HISTORY TABLE (Table 7)
    #     Cols: Date, Author, Version, Change brief

    change_history = data.get("changeHistory", [])
    if change_history:
        _fill_table(tables[7], change_history, 4)

    # -- Write to buffer --
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
