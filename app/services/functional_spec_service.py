
import io
import os
from typing import List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Paths & constants ────────────────────────────────────

TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..", "static", "docs",
    "Functional_Specification_Template.docx",
)

# XML namespace for Word 2010 checkbox SDT elements
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"

BSH_GREY = RGBColor(0x64, 0x74, 0x8B)


# ─── Low-level helpers ────────────────────────────────────

def _get_cell_label(cell) -> str:
    """Extract label text from a checkbox table cell (strips checkbox char and whitespace)."""
    return cell.text.strip().lstrip("\u2610\u2612 ").strip()


def _check_cell_checkbox(cell):
    """Toggle a Word SDT checkbox inside a table cell to checked."""
    for sdt in cell._element.iter(qn("w:sdt")):
        # Set w14:checked val="1"
        for elem in sdt.iter():
            if elem.tag == f"{{{W14_NS}}}checked":
                elem.set(f"{{{W14_NS}}}val", "1")
        # Swap display glyph
        for t_elem in sdt.iter(qn("w:t")):
            if t_elem.text and "\u2610" in t_elem.text:
                t_elem.text = t_elem.text.replace("\u2610", "\u2612")


def _toggle_column_checkboxes(table, col_selections: dict, header_row: bool = True):
    """
    Check table cells per-column.
    col_selections: {col_index: [selected_items]}
    Handles 'Other' appearing in multiple columns without cross-contamination.
    """
    start_row = 1 if header_row else 0
    for col_idx, selected in col_selections.items():
        normalised = {s.strip().lower() for s in selected if s}
        if not normalised:
            continue
        for row in table.rows[start_row:]:
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                label = _get_cell_label(cell)
                if label and label.lower() in normalised:
                    _check_cell_checkbox(cell)


def _toggle_table_checkboxes(table, selected_items: list, header_row: bool = False):
    """Check all cells whose label matches any item in selected_items (flat list)."""
    start_row = 1 if header_row else 0
    normalised = {s.strip().lower() for s in selected_items if s}
    if not normalised:
        return
    for row in table.rows[start_row:]:
        for cell in row.cells:
            label = _get_cell_label(cell)
            if label and label.lower() in normalised:
                _check_cell_checkbox(cell)


def _set_para_text(para, text: str):
    """Replace a paragraph's text, preserving the first run's formatting."""
    if para.runs:
        for run in para.runs[1:]:
            run.text = ""
        para.runs[0].text = text
    else:
        para.add_run(text)


def _remove_paragraph(para):
    """Remove a paragraph element from the document body."""
    p = para._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _insert_text_para_after(ref_element, text: str, bold=False, italic=False):
    """Insert a new w:p with text after ref_element. Returns the new w:p element."""
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")

    rpr = OxmlElement("w:rPr")
    if bold:
        rpr.append(OxmlElement("w:b"))
    if italic:
        rpr.append(OxmlElement("w:i"))
    new_r.append(rpr)

    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = text
    new_r.append(new_t)
    new_p.append(new_r)

    ref_element.addnext(new_p)
    return new_p


def _insert_images_after(doc, ref_element, images: list):
    """Insert images into the document body right after ref_element."""
    if not images:
        return ref_element

    body = doc.element.body
    current_ref = ref_element

    for img_info in images:
        img_stream = io.BytesIO(img_info["data"])
        try:
            # Add picture to end of document (creates relationship + element)
            doc.add_picture(img_stream, width=Inches(3.0))
            pic_para = doc.paragraphs[-1]
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Move the picture element to the correct position
            pic_elem = pic_para._element
            body.remove(pic_elem)
            current_ref.addnext(pic_elem)
            current_ref = pic_elem

            # Add a caption paragraph after the picture
            cap_p = OxmlElement("w:p")
            # Center alignment
            ppr = OxmlElement("w:pPr")
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "center")
            ppr.append(jc)
            cap_p.append(ppr)
            # Caption run (italic, 9pt, grey)
            cap_r = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            rpr.append(OxmlElement("w:i"))
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "18")  # 9pt in half-points
            rpr.append(sz)
            clr = OxmlElement("w:color")
            clr.set(qn("w:val"), "64748B")
            rpr.append(clr)
            cap_r.append(rpr)
            cap_t = OxmlElement("w:t")
            cap_t.text = img_info.get("name", "Screenshot")
            cap_r.append(cap_t)
            cap_p.append(cap_r)

            current_ref.addnext(cap_p)
            current_ref = cap_p

        except Exception:
            err_p = _insert_text_para_after(
                current_ref,
                f"[Image: {img_info.get('name', 'unknown')} — could not be embedded]",
                italic=True,
            )
            current_ref = err_p

    return current_ref


# ─── Main generator ───────────────────────────────────────

def generate_functional_spec_docx(
    data: dict,
    problem_images: List[dict] = None,
    solution_images: List[dict] = None,
) -> io.BytesIO:
    """
    Fill the BSH Functional Specification template with form data.
    Returns an in-memory BytesIO buffer ready for streaming.

    Template paragraph-index map (Functional_Specification_Template.docx):
      [ 0] Heading 1  "Functional Specification"
      [ 1] spacer
      [ 2] Heading 2  "User Story"
      [ 3] Normal     "I as a ..."
      [ 4] Normal     "want to ..."
      [ 5] Normal     "to be able to ..."
      [ 6] spacer
      [ 7] Heading 2  "Process"
      [ 8] spacer                       (before Table 0)
          TABLE 0 -- Process checkboxes  (10 rows x 3 cols)
      [ 9] spacer                       (after Table 0)
      [10] Normal     "If you did not find a suitable Function and Area ..."
      [11] spacer                       (describe-below placeholder)
      [12] Heading 2  "User"
      [13] spacer                       (before Table 1)
          TABLE 1 -- User checkboxes     (3 rows x 3 cols, no header row)
      [14] spacer                       (after Table 1)
      [15] Normal     "If you did not find a suitable User-Group ..."
      [16] spacer                       (describe-below placeholder)
      [17] Heading 2  "Actual Problem what should be solved"
      [18] Normal     "Describe the Problem ..."
      [19-22] spacers                   (content area)
      [23] Heading 2  "Solution Description"
      [24] Normal     "Describe how could the solution ..."
      [25-27] spacers                   (content area)
      [28] Heading 1  "Solution Design"
      [29] spacer
      [30] Heading 2  "Development System"
      [31] spacer                       (before Table 2)
          TABLE 2 -- Dev System checkboxes (7 rows x 3 cols)
      [32] spacer                       (after Table 2)
      [33] Heading 2  "Technical Details"
      [34] Normal     "Describe all objects ..."
      [35-41] spacers                   (content area)
      [42] Heading 2  "Names and Language"
      [43] Normal     "If you create new fields ..."
      [44-45] spacers                   (content area)
      [46] Heading 2  "Authorization"
      [47] Normal     "Consider special authorization objects ..."
      [48-50] spacers                   (content area)
    """
    doc = Document(TEMPLATE_PATH)
    paras = doc.paragraphs  # snapshot -- element refs stay valid after edits

    process_table = doc.tables[0]
    user_table = doc.tables[1]
    dev_table = doc.tables[2]

    # ================================================================
    # 1. USER STORY (paras 3-5)
    # ================================================================
    user_story = data.get("userStory", {})
    role = user_story.get("role", "").strip()
    want = user_story.get("want", "").strip()
    ability = user_story.get("ability", "").strip()

    if role:
        _set_para_text(paras[3], f"I as a {role}")
    if want:
        _set_para_text(paras[4], f"want to {want}")
    if ability:
        _set_para_text(paras[5], f"to be able to {ability}")

    # ================================================================
    # 2. PROCESS -- Table 0 checkboxes (per-column to avoid "Other" clash)
    # ================================================================
    process = data.get("process", {})
    _toggle_column_checkboxes(process_table, {
        0: process.get("function", {}).get("selected", []),
        1: process.get("processArea", {}).get("selected", []),
        2: process.get("processSubArea", {}).get("selected", []),
    }, header_row=True)

    # Remove the instruction hint paragraph [10]
    _remove_paragraph(paras[10])

    # Process custom "Other" text and describe-below
    custom_fn = process.get("function", {}).get("other", "").strip()
    custom_pa = process.get("processArea", {}).get("other", "").strip()
    process_desc = process.get("describeBelow", "").strip()

    has_other = custom_fn or custom_pa or process_desc
    if has_other:
        parts = []
        if custom_fn:
            parts.append(f"Function Other: {custom_fn}")
        if custom_pa:
            parts.append(f"Area Other: {custom_pa}")
        if process_desc:
            parts.append(process_desc)
        _set_para_text(paras[11], "Other: " + "  |  ".join(parts))
    else:
        _remove_paragraph(paras[11])

    # ================================================================
    # 3. USER -- Table 1 checkboxes (flat list, no header row)
    # ================================================================
    user_section = data.get("user", {})
    _toggle_table_checkboxes(user_table, user_section.get("selected", []), header_row=False)

    # Remove the instruction hint paragraph [15]
    _remove_paragraph(paras[15])

    # User custom "Other" text and describe-below
    user_other = user_section.get("other", "").strip()
    user_desc = user_section.get("describeBelow", "").strip()

    has_user_other = user_other or user_desc
    if has_user_other:
        parts = []
        if user_other:
            parts.append(user_other)
        if user_desc:
            parts.append(user_desc)
        _set_para_text(paras[16], "Other: " + "  |  ".join(parts))
    else:
        _remove_paragraph(paras[16])

    # ================================================================
    # 4. PROBLEM DESCRIPTION (para 18 = instruction, 19 = first content slot)
    # ================================================================
    _remove_paragraph(paras[18])  # Remove "Describe the Problem ..." instruction

    problem_text = data.get("problemDescription", "").strip()
    if problem_text:
        _set_para_text(paras[19], problem_text)
        for p in [paras[22], paras[21], paras[20]]:
            _remove_paragraph(p)

    problem_ref = paras[19]._element
    _insert_images_after(doc, problem_ref, problem_images)

    # ================================================================
    # 5. SOLUTION DESCRIPTION (para 24 = instruction, 25 = content slot)
    # ================================================================
    _remove_paragraph(paras[24])  # Remove "Describe how could the solution ..." instruction

    solution_text = data.get("solutionDescription", "").strip()
    if solution_text:
        _set_para_text(paras[25], solution_text)
        for p in [paras[27], paras[26]]:
            _remove_paragraph(p)

    solution_ref = paras[25]._element
    _insert_images_after(doc, solution_ref, solution_images)

    # ================================================================
    # 6. DEVELOPMENT SYSTEM -- Table 2 checkboxes (per-column)
    # ================================================================
    dev_system = data.get("developmentSystem", {})
    _toggle_column_checkboxes(dev_table, {
        0: dev_system.get("erp", {}).get("selected", []),
        1: dev_system.get("scm", {}).get("selected", []),
        2: dev_system.get("cloud", {}).get("selected", []),
    }, header_row=True)

    custom_erp = dev_system.get("erp", {}).get("other", "").strip()
    custom_scm = dev_system.get("scm", {}).get("other", "").strip()
    custom_cloud = dev_system.get("cloud", {}).get("other", "").strip()
    if custom_erp or custom_scm or custom_cloud:
        parts = []
        if custom_erp:
            parts.append(f"ERP: {custom_erp}")
        if custom_scm:
            parts.append(f"SCM: {custom_scm}")
        if custom_cloud:
            parts.append(f"Cloud: {custom_cloud}")
        _set_para_text(paras[32], "Other: " + "  |  ".join(parts))

    # ================================================================
    # 7. TECHNICAL DETAILS (para 34 = instruction, 35 = content slot)
    # ================================================================
    _remove_paragraph(paras[34])  # Remove "Describe all objects ..." instruction

    tech_text = data.get("technicalDetails", "").strip()
    if tech_text:
        _set_para_text(paras[35], tech_text)
        for p in [paras[41], paras[40], paras[39], paras[38], paras[37], paras[36]]:
            _remove_paragraph(p)

    # ================================================================
    # 8. NAMES AND LANGUAGE (para 43 = instruction, 44 = content slot)
    # ================================================================
    _remove_paragraph(paras[43])  # Remove "If you create new fields ..." instruction

    names_text = data.get("namesAndLanguage", "").strip()
    if names_text:
        _set_para_text(paras[44], names_text)
        _remove_paragraph(paras[45])

    # ================================================================
    # 9. AUTHORIZATION (para 47 = instruction, 48 = content slot)
    # ================================================================
    _remove_paragraph(paras[47])  # Remove "Consider special authorization ..." instruction

    auth_text = data.get("authorization", "").strip()
    if auth_text:
        _set_para_text(paras[48], auth_text)
        for p in [paras[50], paras[49]]:
            _remove_paragraph(p)

    # -- Write to buffer --
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
